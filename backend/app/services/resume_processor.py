"""Resume file processing service supporting multiple formats."""
from __future__ import annotations
import os
from pathlib import Path
from typing import BinaryIO
from pypdf import PdfReader
from docx import Document


class ResumeProcessor:
    """Service for processing resume files in various formats."""

    @staticmethod
    async def extract_text_from_file(file_content: bytes, filename: str) -> str:
        """
        Extract text from a resume file.

        Supports: .txt, .md, .pdf, .doc, .docx

        Args:
            file_content: Binary content of the file
            filename: Name of the file (used to determine format)

        Returns:
            Extracted text content
        """
        file_ext = Path(filename).suffix.lower()

        if file_ext in ['.txt', '.md']:
            return ResumeProcessor._extract_text_plain(file_content)

        elif file_ext == '.pdf':
            return ResumeProcessor._extract_text_pdf(file_content)

        elif file_ext in ['.doc', '.docx']:
            return ResumeProcessor._extract_text_docx(file_content)

        else:
            raise ValueError(f"Unsupported file format: {file_ext}")

    @staticmethod
    def _extract_text_plain(file_content: bytes) -> str:
        """Extract text from plain text or markdown files."""
        try:
            # Try UTF-8 first
            return file_content.decode('utf-8')
        except UnicodeDecodeError:
            # Fallback to latin-1
            return file_content.decode('latin-1', errors='ignore')

    @staticmethod
    def _extract_text_pdf(file_content: bytes) -> str:
        """Extract text from PDF files using pypdf."""
        from io import BytesIO

        pdf_stream = BytesIO(file_content)
        reader = PdfReader(pdf_stream)

        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"

        return text.strip()

    @staticmethod
    def _extract_text_docx(file_content: bytes) -> str:
        """Extract text from DOCX files using python-docx."""
        from io import BytesIO

        docx_stream = BytesIO(file_content)
        doc = Document(docx_stream)

        # Extract text from paragraphs
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += "\n" + cell.text

        return text.strip()


# Global service instance
resume_processor = ResumeProcessor()
